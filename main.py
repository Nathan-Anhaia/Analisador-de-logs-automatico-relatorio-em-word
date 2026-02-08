import re
from datetime  import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# Função para extrair os dados do log usando expressão regular
def extrair_dados(log):
    logs = []
    expressao = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}:\d{2}:\d{2})\s*(\w*)\s*(.+)"
    resultados = re.findall(expressao, log)
    if resultados:
        for resultado in resultados:
            logs.append({
                "data": resultado[0],
                "hora": resultado[1],
                "tipo": resultado[2],
                "mensagem": resultado[3]
            })
        return logs
    else:
        print("Nenhum log encontrado.")
        return False

# Função para criar o relatório em formato Word
def criar_relatório():
    with open("logs.txt", encoding="utf-8") as log:
        log_completo = log.read()
        extrair_dados(log_completo)
    logs = extrair_dados(log_completo)
    documento = Document()
    documento.add_heading("Relatório de análise de Logs")
    documento.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Conta a quantidade de logs por tipo
    logs_tipo_qtd = {}
    for log in logs:
        if log['tipo'] in logs_tipo_qtd:
            logs_tipo_qtd[log['tipo']] += 1
        else:
            logs_tipo_qtd[log['tipo']] = 1

    # Adiciona a quantidade de logs por tipo ao documento
    documento.add_paragraph(f"\nQuantidade de logs por tipo: ")
    for tipo, qtd in logs_tipo_qtd.items():
        documento.add_paragraph(f"{tipo}: {qtd}")

    # Adiciona os erros ocorridos ao documento
    documento.add_paragraph(f"\nErros ocorridos: ")
    for log in logs:
        if log['tipo']== "ERROR":
            documento.add_paragraph(f"{log['data']} {log['hora']} - {log['mensagem']}")

    documento.add_page_break()

    # Conta a quantidade de logs por dia e tipo
    registros_por_dia = {}
    for log in logs:
        if log['data'] not in registros_por_dia:
            registros_por_dia[log['data']] = {}
        if log['tipo'] not in registros_por_dia[log['data']]:
            registros_por_dia[log['data']][log['tipo']] = 0
        registros_por_dia[log['data']][log['tipo']] += 1

    # Cria a tabela para mostrar a quantidade de logs por dia e tipo
    tabela = documento.add_table(0, 5, 'Light List Accent 1')
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    linha = tabela.add_row().cells
    linha[0].text = "Data"
    linha[1].text = "ERROR"
    linha[2].text = "INFO"
    linha[3].text = "WARNING"
    linha[4].text = "DEBUG"

    # Adiciona os dados de quantidade de logs por dia e tipo à tabela
    for data in sorted(registros_por_dia):
        linha = tabela.add_row().cells
        linha[0].text = data
        linha[1].text = str(registros_por_dia[data].get("ERROR", 0))
        linha[2].text = str(registros_por_dia[data].get("INFO", 0))
        linha[3].text = str(registros_por_dia[data].get("WARNING", 0))
        linha[4].text = str(registros_por_dia[data].get("DEBUG", 0))
    
    # Centraliza o texto em todas as células da tabela
    for linha in tabela.rows:
        for celula in linha.cells:
            celula.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salva o documento com o nome "relatorio_logs_DD-MM-AAAA.docx"                         
    documento.save(f"relatorio_logs_{datetime.now().strftime('%d-%m-%Y')}.docx")

# Chama a função para criar o relatório
criar_relatório()